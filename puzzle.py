from rdkit import Chem
import numpy as np
import matplotlib.pyplot as plt
from rdkit.Chem import Descriptors


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls



J = 0.01

#pascals triangle function for splitting
def pascal(n):
    row = [1]
    for k in range(n):
        row = [1] + [row[i] + row[i+1] for i in range(len(row)-1)] + [1]
    return row

#chemical shifts
def predict_shift(atom):
    if atom.GetIsAromatic():
        return 7.2
    if atom.GetHybridization().name == "SP3":
        neighbors = [n.GetAtomicNum() for n in atom.GetNeighbors()]

        if 8 in neighbors:
            return 3.5
        return 1.2
    return 2.0

#determining splitting
def multiplicity(atom):
    neighbors = atom.GetNeighbors()
    h_neighbors = sum(n.GetTotalNumHs() for n in neighbors)
    return h_neighbors + 1

#splitting
def split_peaks(center,n_neighbors):
    intensities = pascal(n_neighbors)
    n = len(intensities)
    offsets = np.linspace(-J*(n-1)/2, J*(n-1)/2,n)
    return [(center + offsets[i], intensities[i]) for i in range(n)]
    
#generating spectrum
def gaussian(x, mu, sigma):
    return np.exp(-(x-mu)**2 / (2*sigma**2))



def predict_nmr(molecule):

    #load molecule
    mol = Chem.MolFromSmiles(molecule)

    #grouping equivalent hydrogens
    environments = []

    for atom in mol.GetAtoms():
        if atom.GetAtomicNum() == 6:
            h_count = atom.GetTotalNumHs()
            if h_count > 0:
                environments.append((atom, h_count))
    x = np.linspace(0,10,10000)
    y = np.zeros_like(x)
    for atom, h_count in environments:
        shift = predict_shift(atom)
        n_neighbors = multiplicity(atom) - 1
        peaks = split_peaks(shift, n_neighbors)

        for pos, intensity in peaks:
            y += intensity * h_count *gaussian(x, pos, 0.02)
    
    #plotting nmr spectra
    
    plt.plot(x,y)
    plt.gca().invert_xaxis()
    plt.title("NMR Spectrum")
    plt.xlabel("ppm")
    plt.ylabel("intensity")
    
    plt.savefig("NMR_plot.png", bbox_inches='tight')
    plt.close()


#function to predict mass spectra

def merge_peaks(peaks):
    merged ={}
    for m, inten in peaks:
        key = int(m)
        merged[key] = merged.get(key, 0) + inten
    return list(merged.items())

def normalize(peaks):
    max_i = max(i for _, i in peaks)
    return [(m, i/max_i*100) for m, i in peaks]

def plot_ms(peaks):
    peaks = sorted(peaks)
    for m, i in peaks:
        plt.vlines(m, 0, i)
        
    sorted_peaks = sorted(peaks, key=lambda x: x[0])  # sort by m/z

    min_spacing = 5
    last_m = None
    level = 0

    for m, i in sorted_peaks:
        if last_m is not None and abs(m - last_m) < min_spacing:
            level += 1
        else:
            level = 0

        y_offset = 2 + level * 6  # raise each overlapping label

        plt.text(m, i + y_offset, str(m), ha='center', fontsize=8)

        last_m = m
    plt.title("MS Spectrum")
    plt.xlim(0, max(m for m, _ in peaks) + 10)
    plt.xlabel("m/z")
    plt.ylabel("intensity")
    plt.ylim(0,100)
    plt.savefig("MS_plot.png", bbox_inches='tight')
    plt.close()

def has_mclafferty(mol):
    for atom in mol.GetAtoms():
        if atom.GetAtomicNum() ==6:
            #look for O db
            for bond in atom.GetBonds():
                if bond.GetBondType().name == "DOUBLE":
                    other = bond.GetOtherAtom(atom)
                    if other.GetAtomicNum() == 8:
                        #look for gamma C
                        neighbors = atom.GetNeighbors()
                        for beta in neighbors:
                            for gamma in beta.GetNeighbors():
                                if gamma.GetTotalNumHs() >  0:
                                    return True
    return False

def predict_ms(molecule):
    
    mol = Chem.MolFromSmiles(molecule)

    mw = Descriptors.MolWt(mol)

    peaks = []

    peaks.append((mw, 100)) #set molecular ion to be initial base peak
    num_atoms = mol.GetNumAtoms()

    if num_atoms < 5:
           # restrict fragmentation
        peaks = [(mw, 100)]
    else:

        #creating fragments from single bond cleavages
        for bond in mol.GetBonds():
            bond_idx = bond.GetIdx()

            
            frags = Chem.FragmentOnBonds(mol, [bond.GetIdx()], addDummies=True)

            frag_mols = Chem.GetMolFrags(frags, asMols=True, sanitizeFrags=False)

            clean_frags = []
            for f in frag_mols:
                try:
                    Chem.SanitizeMol(f)
                    clean_frags.append(f)
                except:
                    continue


        
        #common neutral losses
        
        possible_losses = []

        # Only add loss if chemically plausible
        formula = Chem.rdMolDescriptors.CalcMolFormula(mol)

        if "O" in formula:
            possible_losses.append(18)  # water loss

        if "C" in formula:
            possible_losses.append(15)  # methyl loss

        if "C" in formula:
            possible_losses.append(28)  # CO / ethylene

        for loss in possible_losses:
            if mw - loss > 10:
                peaks.append((mw - loss, 30))

        if has_mclafferty(mol):
            peaks.append((mw - 28, 80))
        
        #benzyl/ tropylium ion formation (91m/z)
        if any(atom.GetIsAromatic() for atom in mol.GetAtoms()):
            peaks.append((91,90))
        
        #alpha cleavage to heteroatoms
        if "O" in formula:
            peaks.append((31, 50)) # CH2OH+

        #stable carbocations
        for atom in mol.GetAtoms():
            if atom.GetAtomicNum() ==6 and len(atom.GetNeighbors()) >= 3:
                peaks.append((57, 60))
    
    plot_ms(normalize(merge_peaks(peaks)))
    


def predict_ir(molecule):
    
    patterns = {
        "OH": Chem.MolFromSmarts("[OX2H]"),
        "C=O": Chem.MolFromSmarts("C=O"),
        "NH": Chem.MolFromSmarts("[NX3H]"),
        "aromatic": Chem.MolFromSmarts("c1ccccc1"),
    }

    mol = Chem.MolFromSmiles(molecule)

    peaks = []
    
    if mol.HasSubstructMatch(patterns["OH"]):
        peaks.append((3400, 300, 0.8))  

    if mol.HasSubstructMatch(patterns["C=O"]):
        peaks.append((1700, 50, 1.0))

    if mol.HasSubstructMatch(patterns["NH"]):
        peaks.append((3300, 80, 0.7))

    if mol.HasSubstructMatch(patterns["aromatic"]):
        peaks.append((1600, 50, 0.6))

    
    peaks.append((2900, 80, 0.5))

    x = np.linspace(4000, 500, 8000)  
    y = np.zeros_like(x)

    for center, width, intensity in peaks:
        y += intensity * gaussian(x, center, width)

    y = 1 - y
    
    plt.plot(x, y)
    plt.gca().invert_xaxis()
    plt.title('IR Spectrum')
    plt.xlabel("Wavenumber (cm⁻¹)")
    plt.ylabel("Transmittance")
    plt.savefig("IR_plot.png", bbox_inches='tight')
    plt.close()


def generate_plots(molecule):
    predict_nmr(molecule)
    predict_ms(molecule)
    predict_ir(molecule)


#generate_plots("CCOC(=O)c1ccccc1")


def add_top_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders %s><w:top w:val="single" w:sz="12"/></w:tcBorders>' % nsdecls('w')
    )
    tcPr.append(tcBorders)


def generate_document(smiles_list, output_file="nmr_puzzle.docx"):
    doc = Document()
    problems_per_page = 3

    for i in range(0, len(smiles_list), problems_per_page):

        if i != 0:
            doc.add_page_break()

        
        table = doc.add_table(rows=9, cols=3)

        page_chunk = smiles_list[i:i + problems_per_page]

        for j, smiles in enumerate(page_chunk):

            generate_plots(smiles)

            nmr = "NMR_plot.png"
            ms = "MS_plot.png"
            ir = "IR_plot.png"

            
            row_header = j * 3
            row_top = row_header + 1
            row_bottom = row_header + 2

           
            header_cell = table.cell(row_header, 0)
            header_cell.merge(table.cell(row_header, 2))

            
            add_top_border(header_cell)

            
            paragraph = header_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run(f"Problem {i + j + 1}")
            run.bold = True
            run.font.size = Pt(12)

            
            for col, img in enumerate([nmr, ms, ir]):
                cell = table.cell(row_top, col)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(img, width=Inches(1.6))

            
            for col in range(3):
                table.cell(row_bottom, col).text = "\n" * 6

        
        for row_idx, row in enumerate(table.rows):

            if row_idx % 3 == 0:
                row.height = Inches(0.4)   
            elif row_idx % 3 == 1:
                row.height = Inches(1.8)   
            else:
                row.height = Inches(1.0)   

            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

       
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.space_before = 0
                    paragraph.space_after = 0

    doc.save(output_file)




smiles_list = [
    "CCOC(=O)c1ccccc1",
    "CC(C)COC(=O)C",
    "CC(=O)c1ccccc1",
    "CCC(C)O",
    "c1ccccc1CO",
    "CCOCC"
]


generate_document(smiles_list)